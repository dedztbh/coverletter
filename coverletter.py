from docx import Document
import argparse
import re


# https://stackoverflow.com/a/42829667/5394586
def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


def replace_template(indoc, outdoc, template, name):
    doc = Document(indoc)
    regex = re.compile(template)
    docx_replace_regex(doc, regex, name)
    doc.save(regex.sub(name, outdoc))


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Replace template in document')
    parser.add_argument('name')
    parser.add_argument('--in', help='Input file name', default='in.docx')
    parser.add_argument('--out', help='Output file name', default='<company name> Cover Letter.docx')
    parser.add_argument('--template', help='Template regex', default='<company name>')
    args = parser.parse_args()
    replace_template(args.__getattribute__('in'), args.out, args.template, args.name)
